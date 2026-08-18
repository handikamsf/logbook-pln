import streamlit as st
import json
import base64
import time
import sqlite3
import os
import pandas as pd
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from datetime import datetime
import pytz

# ================= KONFIGURASI HALAMAN =================
st.set_page_config(
    page_title="Sistem Logbook Magang - PT PLN Electricity Services",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ================= DATABASE SQLITE =================
DB_FILE = "logbook_database.db"

def init_db():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS logbook_entries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            minggu_key TEXT NOT NULL,
            tanggal TEXT NOT NULL,
            foto_list TEXT NOT NULL,
            kegiatan TEXT,
            is_libur INTEGER NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def get_all_data():
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("SELECT id, minggu_key, tanggal, foto_list, kegiatan, is_libur FROM logbook_entries ORDER BY id ASC")
    rows = c.fetchall()
    conn.close()
    
    data_dict = {}
    for row in rows:
        row_id, m_key, tgl, f_list, keg, libur = row
        if m_key not in data_dict:
            data_dict[m_key] = []
        data_dict[m_key].append({
            "id": row_id,
            "minggu_key": m_key,
            "tanggal": tgl,
            "foto_b64_list": json.loads(f_list),
            "kegiatan": keg,
            "is_libur": bool(libur)
        })
    return data_dict

def insert_data(minggu_key, tanggal, foto_list, kegiatan, is_libur):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
        INSERT INTO logbook_entries (minggu_key, tanggal, foto_list, kegiatan, is_libur)
        VALUES (?, ?, ?, ?, ?)
    ''', (minggu_key, tanggal, json.dumps(foto_list), kegiatan, int(is_libur)))
    conn.commit()
    conn.close()

def update_data(row_id, minggu_key, kegiatan, is_libur):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute('''
        UPDATE logbook_entries 
        SET minggu_key = ?, kegiatan = ?, is_libur = ?
        WHERE id = ?
    ''', (minggu_key, kegiatan, int(is_libur), row_id))
    conn.commit()
    conn.close()

def delete_data(row_id):
    conn = sqlite3.connect(DB_FILE)
    c = conn.cursor()
    c.execute("DELETE FROM logbook_entries WHERE id = ?", (row_id,))
    conn.commit()
    conn.close()

init_db()

if "edit_id" not in st.session_state:
    st.session_state.edit_id = None

# ================= TEMA UI FORMAL & ENTERPRISE =================
def apply_corporate_theme():
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif !important;
    }
    
    .stApp {
        background-color: #080f1e;
        color: #e2e8f0;
    }

    ::-webkit-scrollbar { width: 8px; height: 8px; }
    ::-webkit-scrollbar-track { background: #0f172a; }
    ::-webkit-scrollbar-thumb { background: #3b82f6; border-radius: 4px; }
    ::-webkit-scrollbar-thumb:hover { background: #60a5fa; }

    [data-testid="stVerticalBlockBorderWrapper"] {
        background: #111c30 !important;
        border: 1px solid #1e293b !important;
        border-radius: 12px !important;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.3) !important;
        transition: all 0.3s ease;
    }
    [data-testid="stVerticalBlockBorderWrapper"]:hover {
        border-color: rgba(56, 189, 248, 0.4) !important;
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.4) !important;
    }

    .stTextInput input, .stTextArea textarea, .stNumberInput input, .stSelectbox div[data-baseweb="select"] {
        background-color: #0d1527 !important;
        border: 1px solid #334155 !important;
        color: #f8fafc !important;
        border-radius: 8px !important;
        font-size: 0.95rem !important;
        padding: 10px 14px !important;
    }
    .stTextInput input:focus, .stTextArea textarea:focus, .stNumberInput input:focus, .stSelectbox div[data-baseweb="select"]:focus-within {
        border-color: #0ea5e9 !important;
        box-shadow: 0 0 0 1px #0ea5e9 !important;
        background-color: rgba(15, 23, 42, 0.95) !important;
    }

    /* THE REAL FIX: Memaksa Preview Kamera 16:9 & ANTI-MIRROR */
    [data-testid="stCameraInput"] video {
        aspect-ratio: 16 / 9 !important;
        object-fit: cover !important;
        border-radius: 8px !important;
        transform: scaleX(1) !important; /* INI YANG MENGHILANGKAN EFEK CERMIN DI LAYAR PREVIEW */
    }

    .stButton > button[kind="primary"] {
        background: linear-gradient(135deg, #0284c7 0%, #2563eb 100%) !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 8px !important;
        font-weight: 600 !important;
        font-size: 0.95rem !important;
        padding: 12px 24px !important;
        letter-spacing: 0.5px !important;
        transition: all 0.2s ease !important;
        box-shadow: 0 4px 6px -1px rgba(2, 132, 199, 0.4) !important;
    }
    .stButton > button[kind="primary"]:hover {
        transform: translateY(-2px);
        box-shadow: 0 10px 15px -3px rgba(2, 132, 199, 0.5) !important;
        filter: brightness(1.1);
    }
    
    .stButton > button[kind="secondary"] {
        background-color: rgba(30, 41, 59, 0.5) !important;
        color: #cbd5e1 !important;
        border: 1px solid #334155 !important;
        border-radius: 8px !important;
        font-size: 0.85rem !important;
        font-weight: 500 !important;
        transition: all 0.2s ease !important;
    }
    .stButton > button[kind="secondary"]:hover {
        background-color: rgba(51, 65, 85, 0.8) !important;
        color: #f8fafc !important;
        border-color: #475569 !important;
    }

    .stTabs [data-baseweb="tab-list"] {
        background-color: transparent;
        border-bottom: 2px solid #1e293b;
        gap: 15px;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 12px 20px !important;
        background-color: transparent;
        border: none !important;
        color: #94a3b8 !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
    }
    .stTabs [aria-selected="true"] {
        border-bottom: 3px solid #38bdf8 !important;
        color: #f8fafc !important;
    }
    .stTabs [aria-selected="true"] p {
        color: #38bdf8 !important;
    }

    .corporate-header {
        border-bottom: 1px solid #1e293b;
        padding-bottom: 20px;
        margin-bottom: 25px;
    }
    .corporate-title {
        color: #f8fafc;
        font-size: 2rem;
        font-weight: 800;
        margin: 0;
        letter-spacing: -0.5px;
    }
    .corporate-subtitle {
        color: #94a3b8;
        font-size: 1rem;
        margin-top: 5px;
        margin-bottom: 0;
    }

    [data-testid="stSidebar"] {
        background-color: #050a15 !important;
        border-right: 1px solid #1e293b !important;
    }

    [data-testid="stMetricValue"] {
        color: #38bdf8 !important;
        font-size: 2.2rem !important;
        font-weight: 800 !important;
        text-shadow: 0px 2px 10px rgba(56, 189, 248, 0.2);
    }
    [data-testid="stMetricLabel"] {
        color: #cbd5e1 !important;
        font-size: 0.9rem !important;
        font-weight: 600 !important;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }

    .corporate-footer {
        text-align: center;
        margin-top: 50px;
        padding-top: 20px;
        border-top: 1px solid #1e293b;
        color: #475569;
        font-size: 0.85rem;
        font-weight: 500;
    }
    </style>
    """, unsafe_allow_html=True)

apply_corporate_theme()

# ================= IMAGE ENGINE =================
def image_to_base64(image_bytes):
    return base64.b64encode(image_bytes).decode('utf-8')

def base64_to_image(b64_string):
    return BytesIO(base64.b64decode(b64_string))

def add_timestamp_watermark(image_bytes, text):
    img = Image.open(BytesIO(image_bytes)).convert("RGBA")
    
    target_ratio = 16.0 / 9.0
    img_width, img_height = img.size
    current_ratio = img_width / img_height
    
    if current_ratio < target_ratio:
        new_height = int(img_width / target_ratio)
        left, top, right, bottom = 0, (img_height - new_height) / 2, img_width, (img_height + new_height) / 2
        img = img.crop((left, top, right, bottom))
    elif current_ratio > target_ratio:
        new_width = int(img_height * target_ratio)
        left, top, right, bottom = (img_width - new_width) / 2, 0, (img_width + new_width) / 2, img_height
        img = img.crop((left, top, right, bottom))
    
    txt_layer = Image.new('RGBA', img.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(txt_layer)
    try: font = ImageFont.truetype("arial.ttf", int(img.width * 0.035)) 
    except: font = ImageFont.load_default()
    
    bbox = draw.textbbox((0, 0), text, font=font)
    text_w, text_h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    x, y = img.width - text_w - 20, img.height - text_h - 20
    
    draw.rectangle([x-10, y-10, x+text_w+10, y+text_h+10], fill=(0, 0, 0, 160))
    draw.text((x, y), text, font=font, fill=(255, 255, 255, 255))
    
    watermarked = Image.alpha_composite(img, txt_layer)
    output = BytesIO()
    watermarked.convert("RGB").save(output, format="JPEG", quality=95)
    return output.getvalue()

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# ================= DATA WAKTU & ZONA WAKTU WIB =================
tz_wib = pytz.timezone('Asia/Jakarta')
hari_indo = ["Senin", "Selasa", "Rabu", "Kamis", "Jum'at", "Sabtu", "Minggu"]
bulan_indo = ["Januari", "Februari", "Maret", "April", "Mei", "Juni", "Juli", "Agustus", "September", "Oktober", "November", "Desember"]

HOLIDAYS_2026 = {
    "2026-08-17": "Libur Nasional Hari Ulang Tahun (HUT) Proklamasi Kemerdekaan Republik Indonesia",
    "2026-09-16": "Libur Nasional Peringatan Maulid Nabi Muhammad SAW",
    "2026-12-25": "Libur Nasional Hari Raya Natal"
}

# ================= SIDEBAR UI =================
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/9/97/Logo_PLN.png", width=100)
    st.markdown("<h2 style='margin-top: 15px;'>⚙️ Control Panel</h2>", unsafe_allow_html=True)
    
    waktu_wib_sidebar = datetime.now(tz_wib)
    st.caption(f"🕒 Waktu Sistem: {waktu_wib_sidebar.strftime('%d %b %Y - %H:%M WIB')}")
    
    with st.container(border=True):
        st.markdown("**Identitas Mahasiswa**")
        nama_doc = st.text_input("Nama Lengkap", value="", placeholder="Contoh: Budi Santoso")
        npm_doc = st.text_input("NPM / NIM", value="", placeholder="Contoh: 24083010001")
        prodi_doc = st.text_input("Program Studi", value="Sains Data")
        dosen_doc = st.text_input("Dosen Pembimbing", value="", placeholder="Nama Dosen beserta Gelar")
    
    st.markdown("---")
    st.markdown("#### 🛡️ Database Management")
    st.caption("Pusat pencadangan dan pemulihan SQL.")
    
    if os.path.exists(DB_FILE):
        with open(DB_FILE, "rb") as f:
            db_bytes = f.read()
        st.download_button(
            label="⬇️ Backup Database (.db)",
            data=db_bytes,
            file_name=f"Logbook_DB_Backup_{datetime.now(tz_wib).strftime('%Y%m%d')}.db",
            mime="application/octet-stream",
            use_container_width=True
        )
    
    uploaded_db = st.file_uploader("📂 Restore Database", type=['db'])
    if uploaded_db:
        if st.button("⚠️ Jalankan Restore Data", use_container_width=True):
            with open(DB_FILE, "wb") as f:
                f.write(uploaded_db.read())
            st.success("Database dipulihkan! Halaman akan dimuat ulang.")
            time.sleep(1.5)
            st.rerun()

# ================= MAIN HEADER UI =================
st.markdown("""
<div class="corporate-header">
    <h1 class="corporate-title">Workstation Logbook Magang</h1>
    <p class="corporate-subtitle">Platform Pencatatan Harian Terintegrasi — PT PLN Electricity Services Jawa Timur</p>
</div>
""", unsafe_allow_html=True)

# ================= TAB NAVIGASI =================
tab1, tab2, tab3 = st.tabs(["📝 Form Input Kegiatan", "📊 Riwayat & Analytics", "📄 Ekspor Laporan"])

# ----------------- TAB 1: INPUT DATA -----------------
with tab1:
    with st.container(border=True):
        col1, col2 = st.columns(2)

        with col1:
            tanggal_input = st.date_input("Kalender Pelaksanaan", value=datetime.now(tz_wib).date())
            
            with st.container(border=True):
                st.markdown("<p style='font-size: 0.9rem; font-weight: 600; margin-bottom: 5px; color:#cbd5e1;'>⏰ Waktu Sinkronisasi (Watermark)</p>", unsafe_allow_html=True)
                use_current_time = st.toggle("Otomatis ikuti jam saat ini", value=True)
                
                waktu_wib_form = datetime.now(tz_wib)
                
                if use_current_time:
                    waktu_preview = waktu_wib_form.strftime("%H:%M WIB")
                    waktu_str = "AUTO"
                else:
                    col_j, col_m = st.columns(2)
                    with col_j: jam = st.selectbox("Jam", options=[f"{i:02d}" for i in range(24)], index=waktu_wib_form.hour)
                    with col_m: menit = st.selectbox("Menit", options=[f"{i:02d}" for i in range(60)], index=waktu_wib_form.minute)
                    waktu_preview = f"{jam}:{menit} WIB"
                    waktu_str = waktu_preview
            
            START_DATE = datetime(2026, 8, 3).date()
            delta_days = (tanggal_input - START_DATE).days
            auto_minggu = (delta_days // 7) + 1 if delta_days >= 0 else 1
            
            minggu_ke = st.number_input("Minggu Ke-", min_value=1, max_value=50, value=auto_minggu, step=1)
            
            hari_input = hari_indo[tanggal_input.weekday()]
            tanggal_format = f"{tanggal_input.day} {bulan_indo[tanggal_input.month-1]} {tanggal_input.year}"
            tanggal_lengkap = f"{hari_input}, {tanggal_format}"
            
            date_str = tanggal_input.strftime("%Y-%m-%d")
            is_holiday_auto = date_str in HOLIDAYS_2026
            holiday_text = HOLIDAYS_2026.get(date_str, "Libur Nasional / Tanggal Merah")

            st.success(f"📌 Label Cetak: **{tanggal_lengkap} | {waktu_preview}**")
            
            is_libur = st.toggle("🚩 Deklarasikan sebagai Hari Libur / Tanggal Merah", value=is_holiday_auto)
            
            if is_libur:
                kegiatan = st.text_area("Keterangan Hari Libur", value=holiday_text if is_holiday_auto else "", height=140)
                kegiatan_final = kegiatan
            else:
                kategori_tugas = st.selectbox("Kategori Pekerjaan Utama", ["Data Science / Analitik", "Pengembangan Sistem / IT", "Administrasi & Pelaporan", "Rapat / Koordinasi Tim", "Lainnya (Tulis Manual)"])
                kegiatan_manual = st.text_area("Rincian Aktivitas", height=140, placeholder="Jelaskan detail tugas yang dikerjakan hari ini...")
                
                if kategori_tugas != "Lainnya (Tulis Manual)":
                    kegiatan_final = f"[{kategori_tugas}]\n{kegiatan_manual}" if kegiatan_manual else f"[{kategori_tugas}]"
                else:
                    kegiatan_final = kegiatan_manual

        with col2:
            with st.container(border=True):
                st.markdown("<p style='font-size: 0.95rem; font-weight: 600; margin-bottom: 10px; color:#cbd5e1;'>📸 Lampiran Visual Pekerjaan</p>", unsafe_allow_html=True)
                st.markdown("<p style='font-size: 0.85rem; color:#94a3b8; margin-bottom: 15px;'>Kombinasikan foto kamera (Selfie) dan unggah file (Screenshot) secara bersamaan jika dibutuhkan.</p>", unsafe_allow_html=True)
                
                foto_bytes_list = []
                
                st.markdown("**1. Ambil Foto Kamera (Otomatis 16:9 & Anti-Mirror)**")
                aktifkan_kamera = st.toggle("📷 Nyalakan Kamera", value=False)
                
                if aktifkan_kamera:
                    cam_col1, cam_col2 = st.columns(2)
                    with cam_col1:
                        # Hapus ImageOps, langsung simpan gambarnya secara natural
                        foto_kamera_1 = st.camera_input("Foto Kamera 1", key="cam1")
                        if foto_kamera_1:
                            foto_bytes_list.append(foto_kamera_1.getvalue())
                    with cam_col2:
                        foto_kamera_2 = st.camera_input("Foto Kamera 2", key="cam2")
                        if foto_kamera_2:
                            foto_bytes_list.append(foto_kamera_2.getvalue())
                else:
                    st.info("Kamera dinonaktifkan untuk menghemat daya. Nyalakan sakelar di atas untuk mulai berswafoto.")
                
                st.markdown("<hr style='margin: 15px 0; border-color: #1e293b;'>", unsafe_allow_html=True)
                
                st.markdown("**2. Unggah Screenshot / File Tambahan**")
                file_uploads = st.file_uploader("Tarik & Lepas File Visual Di Sini", type=['jpg', 'jpeg', 'png'], accept_multiple_files=True)
                if file_uploads: 
                    for f in file_uploads:
                        foto_bytes_list.append(f.read())
                
                if is_libur:
                    st.info("ℹ️ Dokumen visual tidak diwajibkan untuk hari libur.")
                else:
                    st.caption("✨ **Keterangan:** Layar kamera kini tidak seperti cermin lagi, sehingga teks apapun akan terbaca normal sebelum dan sesudah dijepret.")

        st.markdown("<br>", unsafe_allow_html=True)
        
        if st.button("🚀 PROSES & SIMPAN KE DATABASE", type="primary", use_container_width=True):
            if not is_libur and not foto_bytes_list:
                st.error("Transaksi Ditolak: Logbook hari kerja wajib menyertakan minimal 1 bukti visual (dari kamera ATAU file upload).")
            elif not is_libur and not kegiatan_manual.strip() and kategori_tugas == "Lainnya (Tulis Manual)":
                 st.error("Transaksi Ditolak: Rincian aktivitas wajib diisi.")
            else:
                minggu_key = f"Minggu ke-{minggu_ke}"
                
                if waktu_str == "AUTO":
                    waktu_aktual = datetime.now(tz_wib).strftime("%H:%M WIB")
                    timestamp_final = f"{tanggal_lengkap} | {waktu_aktual}"
                else:
                    timestamp_final = f"{tanggal_lengkap} | {waktu_str}"
                
                foto_b64_list = []
                if foto_bytes_list:
                    for f_bytes in foto_bytes_list:
                        watermarked_bytes = add_timestamp_watermark(f_bytes, timestamp_final)
                        foto_b64_list.append(image_to_base64(watermarked_bytes))
                
                insert_data(minggu_key, tanggal_lengkap, foto_b64_list, kegiatan_final.strip(), is_libur)
                
                st.balloons()
                msg = st.success(f"✅ Transaksi Berhasil! Logbook tanggal {tanggal_lengkap} telah dienkripsi ke database.")
                time.sleep(1.5)
                msg.empty()
                st.rerun()

# ----------------- TAB 2: REVIEW DATA & ANALYTICS -----------------
with tab2:
    db_data = get_all_data()
    
    if not db_data:
        st.info("Repositori kosong. Silakan input data kegiatan terlebih dahulu.")
    else:
        tot_minggu = len(db_data)
        tot_kegiatan = sum(len(keg) for keg in db_data.values())
        tot_libur = sum(1 for items in db_data.values() for item in items if item['is_libur'])
        tot_kerja = tot_kegiatan - tot_libur
        
        total_jam_kerja = tot_kerja * 9.5
        progress_val = min(tot_minggu / 20.0, 1.0)
        
        with st.container(border=True):
            st.markdown("#### 📊 Ringkasan Eksekutif & Jam Kerja")
            m_col1, m_col2, m_col3, m_col4 = st.columns(4)
            m_col1.metric("Periode Aktif", f"{tot_minggu} Minggu")
            m_col2.metric("Hari Kerja Efektif", f"{tot_kerja} Hari")
            m_col3.metric("Akumulasi Jam Kerja", f"{total_jam_kerja:g} Jam", help="Dihitung dari 07.30 s.d. 17.00 (9,5 Jam per hari efektif)")
            m_col4.metric("Izin / Libur", f"{tot_libur} Hari")
            st.progress(progress_val, text=f"Estimasi Penyelesaian Program: {int(progress_val*100)}% (Berdasarkan target 20 Minggu)")
        
        st.markdown("<br>", unsafe_allow_html=True)

        with st.container(border=True):
            st.markdown("#### 📈 Analisis Produktivitas Mingguan")
            minggu_labels = []
            jumlah_aktivitas = []
            
            for m_key, items in sorted(db_data.items(), key=lambda x: int(x[0].split('-')[1])):
                minggu_labels.append(m_key)
                aktif = sum(1 for i in items if not i['is_libur'])
                jumlah_aktivitas.append(aktif)
            
            if minggu_labels:
                df_chart = pd.DataFrame({
                    "Minggu": minggu_labels,
                    "Jumlah Hari Kerja": jumlah_aktivitas
                })
                st.bar_chart(df_chart.set_index("Minggu"), color="#0ea5e9", use_container_width=True)

        st.markdown("<br>", unsafe_allow_html=True)
        
        st.markdown("#### 🔍 Penjelajah Data (Database Explorer)")
        f_col1, f_col2 = st.columns([1, 2])
        with f_col1:
            pilihan_minggu = ["Semua Periode"] + sorted(list(db_data.keys()), key=lambda x: int(x.split('-')[1]))
            filter_minggu = st.selectbox("Saring Berdasarkan Minggu", pilihan_minggu)
        with f_col2:
            search_query = st.text_input("Pencarian Spesifik (Cari kegiatan, tanggal, atau kategori...)")
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        filtered_data = {}
        for m_key, items in db_data.items():
            if filter_minggu != "Semua Periode" and m_key != filter_minggu:
                continue
            
            matched_items = []
            for item in items:
                if search_query.lower() in item['kegiatan'].lower() or search_query.lower() in item['tanggal'].lower():
                    matched_items.append(item)
            
            if matched_items:
                filtered_data[m_key] = matched_items
        
        if not filtered_data:
            st.warning("Pencarian tidak menemukan hasil yang cocok di database.")
        else:
            for minggu_key, kegiatan_list in sorted(filtered_data.items(), key=lambda x: int(x[0].split('-')[1])):
                with st.expander(f"📂 {minggu_key} — ({len(kegiatan_list)} Rekaman Data)", expanded=False):
                    for i, item in enumerate(kegiatan_list):
                        
                        if st.session_state.edit_id == item['id']:
                            with st.container(border=True):
                                st.markdown(f"<p style='color:#38bdf8; font-weight:700; font-size:1rem; margin-bottom:10px;'>✏️ Mode Edit Data: {item['tanggal']}</p>", unsafe_allow_html=True)
                                
                                edit_minggu = st.text_input("Grup Minggu", value=item['minggu_key'], key=f"edit_m_{item['id']}")
                                edit_libur = st.toggle("Tandai Libur", value=item['is_libur'], key=f"edit_l_{item['id']}")
                                edit_kegiatan = st.text_area("Deskripsi Pekerjaan / Keterangan Libur", value=item['kegiatan'], key=f"edit_k_{item['id']}", height=120)
                                
                                st.caption("*(Keterangan: Modifikasi dokumen visual/foto tidak diizinkan pada Mode Edit untuk menjaga integritas rasio & watermark database).*")
                                
                                col_sv, col_cx = st.columns(2)
                                with col_sv:
                                    if st.button("💾 Simpan Pembaruan", key=f"save_{item['id']}", type="primary", use_container_width=True):
                                        update_data(item['id'], edit_minggu, edit_kegiatan, edit_libur)
                                        st.session_state.edit_id = None
                                        st.toast("✅ Perubahan berhasil diamankan ke Database.")
                                        time.sleep(0.5)
                                        st.rerun()
                                    
                                with col_cx:
                                    if st.button("❌ Batal Edit", key=f"cancel_{item['id']}", use_container_width=True):
                                        st.session_state.edit_id = None
                                        st.rerun()
                        else:
                            is_holiday = item.get("is_libur", False)
                            with st.container(border=True):
                                if is_holiday:
                                    st.markdown("<p style='color:#ef4444; font-weight:700; font-size:0.85rem; margin:0;'>🔴 STATUS: HARI LIBUR NASIONAL / IZIN</p>", unsafe_allow_html=True)
                                    
                                prev_col1, prev_col2, prev_col3 = st.columns([1.5, 2.5, 0.7])
                                with prev_col1: 
                                    st.caption(f"**{item['tanggal']}**")
                                    fotos = item.get("foto_b64_list", [])
                                    if fotos:
                                        for idx_img in range(0, len(fotos), 2):
                                            img_cols = st.columns(2)
                                            with img_cols[0]:
                                                st.image(base64_to_image(fotos[idx_img]), use_container_width=True)
                                            if idx_img + 1 < len(fotos):
                                                with img_cols[1]:
                                                    st.image(base64_to_image(fotos[idx_img+1]), use_container_width=True)
                                    else:
                                        st.caption("*(Dokumen visual tidak tersedia)*")
                                with prev_col2:
                                    st.markdown("**Laporan Aktivitas:**")
                                    st.write(item["kegiatan"] if item["kegiatan"] else "-")
                                with prev_col3:
                                    st.write("")
                                    if st.button("✏️ Edit", key=f"edit_sql_{item['id']}", use_container_width=True):
                                        st.session_state.edit_id = item['id']
                                        st.rerun()
                                    if st.button("🗑️ Hapus", key=f"del_sql_{item['id']}", use_container_width=True):
                                        delete_data(item['id'])
                                        st.toast(f"Data Log ID#{item['id']} dieksekusi hapus.")
                                        time.sleep(0.5)
                                        st.rerun()

# ----------------- TAB 3: EXPORT WORD & EXCEL -----------------
with tab3:
    st.markdown("Pusat kompilasi akhir untuk merender data SQL menjadi dokumen laporan resmi.")
    
    col_export1, col_export2 = st.columns(2)
    
    with col_export1:
        with st.container(border=True):
            st.markdown("### 📄 Ekspor Dokumen Resmi (Word)")
            st.caption("Menghasilkan file `.docx` lengkap dengan tabel, foto rasio 16:9, watermark, dan format laporan resmi PLN.")
            
            if st.button("Kompilasi Laporan (.docx)", type="primary", use_container_width=True):
                db_data = get_all_data()
                doc = Document()
                for section in doc.sections:
                    section.top_margin, section.bottom_margin = Inches(1), Inches(1)
                    section.left_margin, section.right_margin = Inches(1), Inches(1)

                p_title = doc.add_paragraph()
                p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_title = p_title.add_run("CATATAN HARIAN/LOGBOOK\nMagang Mandiri PT. PLN Electricity Services Jawa Timur\n(Semester: V TA 2026/2027)")
                run_title.font.bold, run_title.font.name, run_title.font.size = True, 'Times New Roman', Pt(12)
                doc.add_paragraph() 

                meta_data = [
                    ("Nama", nama_doc), ("NPM", npm_doc), ("Program Studi/Jurusan", prodi_doc),
                    ("Dosen Pembimbing", dosen_doc), ("Mitra Satuan Pendidikan", "PT. PLN Electricity Services Jawa Timur"),
                    ("Waktu Pelaksanaan", "3 Agustus - 25 Desember 2026")
                ]
                table_meta = doc.add_table(rows=6, cols=3)
                table_meta.autofit = False
                for idx, width in enumerate([Inches(2.0), Inches(0.2), Inches(4.3)]): 
                    table_meta.columns[idx].width = width

                for i, (key, val) in enumerate(meta_data):
                    row = table_meta.rows[i].cells
                    row[0].text, row[1].text, row[2].text = key, ":", val
                    for cell in row:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            for run in paragraph.runs:
                                run.font.name, run.font.size = 'Times New Roman', Pt(11)
                    row[0].width, row[1].width, row[2].width = Inches(2.0), Inches(0.2), Inches(4.3)

                doc.add_paragraph() 

                for minggu_key, kegiatan_list in sorted(db_data.items(), key=lambda x: int(x[0].split('-')[1])):
                    
                    table = doc.add_table(rows=2, cols=3)
                    table.style = 'Table Grid'
                    table.autofit = False 
                    
                    col_widths = [Inches(0.4), Inches(1.5), Inches(4.6)]
                    for j, col in enumerate(table.columns): col.width = col_widths[j]
                    for row in table.rows:
                        for idx, width in enumerate(col_widths): row.cells[idx].width = width
                    
                    title_cells = table.rows[0].cells
                    title_cells[0].merge(title_cells[1]).merge(title_cells[2])
                    title_p = title_cells[0].paragraphs[0]
                    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    title_run = title_p.add_run(minggu_key)
                    title_run.font.bold, title_run.font.name = True, 'Times New Roman'
                    
                    hdr_cells = table.rows[1].cells
                    hdr_cells[0].text, hdr_cells[1].text, hdr_cells[2].text = 'No.', 'Hari, Tanggal', 'Kegiatan'
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs: run.font.bold, run.font.name = True, 'Times New Roman'

                    for i, item in enumerate(kegiatan_list):
                        row_cells = table.add_row().cells
                        for idx, width in enumerate(col_widths): row_cells[idx].width = width
                            
                        is_libur = item.get("is_libur", False)
                        if is_libur:
                            for cell in row_cells: set_cell_background(cell, "FFE6E6")
                                
                        row_cells[0].text = f"{i+1}."
                        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[1].text = item["tanggal"]
                        
                        keg_cell = row_cells[2]
                        if item["kegiatan"]:
                            keg_cell.paragraphs[0].text = item["kegiatan"]
                            p_img = keg_cell.add_paragraph()
                        else:
                            p_img = keg_cell.paragraphs[0]
                        
                        fotos = item.get("foto_b64_list", [])
                        if fotos:
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for idx_img, img_b64 in enumerate(fotos):
                                run_img = p_img.add_run()
                                run_img.add_picture(base64_to_image(img_b64), width=Inches(2.15)) 
                                
                                if idx_img < len(fotos) - 1:
                                    if (idx_img + 1) % 2 == 0:
                                        p_img = keg_cell.add_paragraph()
                                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        run_img.add_text("   ")
                        
                        for cell in row_cells:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs: run.font.name, run.font.size = 'Times New Roman', Pt(11)

                    doc.add_paragraph() 
                    
                # PENGESAHAN
                teks_nama = nama_doc if nama_doc.strip() else ".........................."
                teks_dosen = dosen_doc if dosen_doc.strip() else ".........................."
                
                table_pengesahan = doc.add_table(rows=2, cols=2)
                table_pengesahan.style = 'Table Grid'
                table_pengesahan.autofit = False
                for row in table_pengesahan.rows:
                    row.cells[0].width, row.cells[1].width = Inches(3.25), Inches(3.25)
                
                table_pengesahan.cell(0, 0).merge(table_pengesahan.cell(0, 1))
                p_peng = table_pengesahan.cell(0, 0).paragraphs[0]
                run_peng = p_peng.add_run("PENGESAHAN")
                run_peng.font.bold, run_peng.font.name = True, 'Times New Roman'
                
                p_dosen = table_pengesahan.cell(1, 0).paragraphs[0]
                p_dosen.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_dosen = p_dosen.add_run(f"Dosen Pembimbing\nMBKM Program Studi\n\n\n\n({teks_dosen})")
                run_dosen.font.name = 'Times New Roman'

                p_mhs = table_pengesahan.cell(1, 1).paragraphs[0]
                p_mhs.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_mhs = p_mhs.add_run(f"Mahasiswa\n\n\n\n\n({teks_nama})")
                run_mhs.font.name = 'Times New Roman'

                doc.add_paragraph() 

                bio = BytesIO()
                doc.save(bio)
                
                file_name_output = f"Logbook_{npm_doc if npm_doc.strip() else 'Magang'}_PLN_Electricity_Services.docx"
                
                st.success("Sintesis dokumen berhasil!")
                st.download_button(
                    label="⬇️ Unduh Dokumen (.docx)",
                    data=bio.getvalue(),
                    file_name=file_name_output,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

    with col_export2:
        with st.container(border=True):
            st.markdown("### 📊 Ekspor Rekap Data (CSV)")
            st.caption("Menghasilkan tabel mentah berisi teks riwayat kegiatan untuk pelaporan cepat di Excel/Spreadsheet.")
            
            db_data = get_all_data()
            if not db_data:
                st.warning("Data masih kosong.")
            else:
                if st.button("Buat Rekapan Excel (.csv)", type="primary", use_container_width=True):
                    csv_data = []
                    for m_key, items in sorted(db_data.items(), key=lambda x: int(x[0].split('-')[1])):
                        for item in items:
                            csv_data.append({
                                "Minggu": m_key,
                                "Tanggal": item['tanggal'],
                                "Status": "Libur / Izin" if item['is_libur'] else "Kerja Aktif",
                                "Kegiatan": item['kegiatan'].replace('\n', ' ')
                            })
                    
                    df_export = pd.DataFrame(csv_data)
                    csv_string = df_export.to_csv(index=False).encode('utf-8')
                    
                    file_name_csv = f"Rekap_Kegiatan_{npm_doc if npm_doc.strip() else 'Magang'}.csv"
                    
                    st.success("Rekap CSV berhasil dibuat!")
                    st.download_button(
                        label="⬇️ Unduh Tabel Laporan (.csv)",
                        data=csv_string,
                        file_name=file_name_csv,
                        mime="text/csv",
                        use_container_width=True
                    )

# ================= FOOTER =================
st.markdown("<div class='corporate-footer'>Dirancang dan Dikembangkan oleh Handika | Enterprise System PLN Electricity Services Jawa Timur 2026</div>", unsafe_allow_html=True)